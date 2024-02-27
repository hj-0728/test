import io
from collections import defaultdict
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from infra_basic.errors import BusinessError
from infra_utility.algorithm.tree import list_to_tree
from infra_utility.datetime_helper import datetime_format, local_now
from infra_utility.file_helper import build_abs_path_by_file
from infra_utility.token_helper import generate_uuid_id

from edu_binshi.model.view.evaluation_assignment_plan_vm import EvaluationAssignmentPlanViewModel
from edu_binshi.model.view.report_indicator_score_vm import ReportIndicatorScoreViewModel
from edu_binshi.model.view.report_table_vm import (
    EnumReportStyle,
    ReportCellStyle,
    ReportTableCellData,
)
from edu_binshi.repository.report_repository_v2 import ReportRepositoryV2
from edu_binshi.utility.word_to_pdf_helper import doc2pdf

ORDER = ["自评", "他评", "综合", "等级"]


class ReportServiceV2:
    def __init__(self, report_repository_v2: ReportRepositoryV2):
        self.__report_repository_v2 = report_repository_v2

        # 从数据库捞出来赋值之后,肯定有值,这里写Optional和默认值,只是为了解决一些波浪线警告
        self.__plan: Optional[EvaluationAssignmentPlanViewModel] = None
        self.__tag_level: int = 1

        # 只定义一个初始值，后面捞回来数据的时候会重新赋值
        self.__max_score_col = 1
        self.__max_indicator_col = 1

    def generate_report(self, evaluation_assignment_id: str):
        """
        生成报告
        """

        document = self.draw_report_doc(evaluation_assignment_id=evaluation_assignment_id)
        doc_name = f"{datetime_format(input_time=local_now(), pattern='%Y%m%d%H%M%S')}_{generate_uuid_id()}"

        stream = io.BytesIO()
        # 保存文档到这个内存文件中
        document.save(stream)
        # 将指针移到流的开头，以便您可以从头开始读取它
        stream.seek(0)
        files = {
            "file": (
                f"{doc_name}.docx",
                stream,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        pdf_data = doc2pdf(files=files)

        # 再次读取 stream 的内容
        stream.seek(0)  # 重置指针
        doc_data = stream.read()

        # 如果需要把文件保存到本地，可以把下面的代码放出来
        # file_path = build_abs_path_by_file(__file__, f"../docs/report/{doc_name}")
        # document.save(f"{file_path}.docx")
        # with open(f"{file_path}.pdf", "wb") as f:
        #     f.write(pdf_data)

        return doc_data, pdf_data

    def draw_report_doc(self, evaluation_assignment_id: str) -> Document:
        """
        绘制表格文件
        """
        plan = self.__report_repository_v2.fetch_evaluation_assignment_plan(
            evaluation_assignment_id=evaluation_assignment_id
        )
        if not plan:
            raise BusinessError(f"评价分配【{evaluation_assignment_id}】的计划不存在")
        criteria_tag = self.__report_repository_v2.fetch_evaluation_criteria_tag(
            evaluation_criteria_id=plan.evaluation_criteria_id,
            executed_finish_at=plan.executed_finish_at,
        )
        if not criteria_tag.tag_level:
            raise BusinessError("评价标准没有打标签")
        self.__plan = plan
        document = Document()
        # 准备报告的内容
        self.add_report_title(document)
        self.add_report_people_info(document)
        self.__tag_level = criteria_tag.tag_level
        if criteria_tag.tag_level == 1:
            # 如果标签打在第一层
            self.add_report_table_for_tag_in_level1(document)
        else:
            self.add_report_table_for_tag_in_level2(document)
        self.add_report_comments(document)

        return document

    def add_report_title(self, document: Document):
        """
        添加报告标题
        """
        style = self._create_report_basic_info_style(styles=document.styles, name="ReportTitle", font_size=15)
        title = document.add_paragraph(text=self.__plan.criteria_name, style=style)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_report_people_info(self, document: Document):
        """
        添加报告的人员信息，如姓名、班级等
        """
        style = self._create_report_basic_info_style(styles=document.styles, name="PeopleInfo", font_size=10.5)
        people_info = f"班级（{self.__plan.dept_name}）    姓名（{self.__plan.people_name}）"
        title = document.add_paragraph(text=people_info, style=style)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def add_report_comments(self, document: Document):
        """
        添加报告的评价标准的备注
        """
        if not self.__plan.criteria_comments:
            return
        style = self._create_report_basic_info_style(styles=document.styles, name="CommentsInfo", font_size=8.5)
        title = document.add_paragraph(text=self.__plan.criteria_comments, style=style)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title.paragraph_format.space_before = Pt(12)

    @staticmethod
    def _create_report_basic_info_style(styles, name: str, font_size: float):
        """
        创建报告基础信息的样式
        """
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor(0, 0, 0)  # black
        style.font.name = "仿宋"
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style._element.attrib[qn("w:styleId")] = name
        return style

    def add_report_table_for_tag_in_level1(self, document: Document):
        """
        添加报告的表格部分
        对标签打在第一层的情况
        """
        row = 0
        table_info: Dict[Tuple, ReportTableCellData] = {}
        merge_info = {}
        tag_tree = self.build_report_indicator_tree()
        for tag in ["N（常量）", "X（变量）"]:
            report_tree_data = tag_tree[tag]
            table_info[row, 0] = ReportTableCellData(value=tag, style=EnumReportStyle.TAG_CELL.name)
            merge_info[(row, 0)] = (row, self.__max_indicator_col + self.__max_score_col - 1)
            row += 1
            row = self.traverse_data_in_level1(
                data_list=report_tree_data,
                table_info=table_info,
                merge_info=merge_info,
                start_row=row,
            )
        self.prepare_merge_data(table_info=table_info, merge_info=merge_info)
        self.create_table_in_word(doc=document, table_info=table_info, merge_info=merge_info)

    def build_report_indicator_tree(self) -> Dict[str, List[ReportIndicatorScoreViewModel]]:
        """
        构建报告的指标树
        """
        data_list = self.__report_repository_v2.fetch_report_indicator_data(
            evaluation_criteria_id=self.__plan.evaluation_criteria_id,
            evaluation_assignment_id=self.__plan.evaluation_assignment_id,
            executed_finish_at=self.__plan.executed_finish_at,
        )
        tag_data_dict = defaultdict(list)
        tag_tree_dict = defaultdict(list)
        max_score_col = 1
        max_indicator_col = 1
        for data in data_list:
            if self.__tag_level == 2 and data.level == 2:
                # 因为标签打在第一层和第二层，对打在第二层的情况，需要把parent_indicator_id设置为None，这样才能构造出来一棵树
                data.parent_indicator_id = None
            if self.__tag_level == 1 and data.level == 1:
                # 等级处是老师自己填，老师自己填的tag是他评，这里写死转换下
                for benchmark in data.benchmark_list:
                    if benchmark.tag == "他评":
                        benchmark.tag = "等级"
            data.sort_benchmark_list()
            tag_data_dict[data.tag].append(data)
            max_score_col = max(max_score_col, len(data.benchmark_list))
            if data.benchmark_list:
                # 如果有benchmark，同层的指标只合并到这个位置
                max_indicator_col = max(max_indicator_col, data.level - 1)
        for tag in ["N（常量）", "X（变量）"]:
            tree = list_to_tree(
                original_list=tag_data_dict[tag],
                tree_node_type=ReportIndicatorScoreViewModel,
                id_attr="indicator_id",
                parent_id_attr="parent_indicator_id",
            )
            tag_tree_dict[tag] = tree
        self.__max_indicator_col = max_indicator_col
        self.__max_score_col = max_score_col
        return tag_tree_dict

    def traverse_data_in_level1(
        self,
        data_list: List[ReportIndicatorScoreViewModel],
        table_info: Dict,
        merge_info: Dict,
        start_row,
    ):
        """
        转换标签打在第一层的数据
        """
        for idx, data in enumerate(data_list, start=1):
            if self.__max_indicator_col == 1:
                start_row = self.traverse_all_data_is_root_in_level1(
                    data=data,
                    table_info=table_info,
                    merge_info=merge_info,
                    first=idx == 1,
                    start_row=start_row,
                )
            else:
                start_row = self.traverse_root_data_in_level1(
                    data=data,
                    table_info=table_info,
                    merge_info=merge_info,
                    first=idx == 1,
                    start_row=start_row,
                )
            start_row = self.traverse_child_data(
                data.children, table_info, merge_info, start_row, 0
            )
        return start_row

    def traverse_all_data_is_root_in_level1(
        self,
        data: ReportIndicatorScoreViewModel,
        table_info: Dict,
        merge_info: Dict,
        first: bool,
        start_row,
    ) -> int:
        """
        处理所有数据都是根节点的情况
        """
        plus = 0
        score_cell_style = EnumReportStyle.SCORE_CELL.name
        if first:
            tag_methods_map = {
                "N（常量）": "观察记录",
                "X（变量）": "观察记录 成果展示",
            }
            table_info[start_row, 0] = ReportTableCellData(
                value="评价内容",
                style=EnumReportStyle.EVALUATION_CONTENT_CELL.name,
            )
            table_info[start_row, 1] = ReportTableCellData(value="评价方式", style=score_cell_style)
            table_info[start_row + 1, 1] = ReportTableCellData(
                value=tag_methods_map[data.tag], style=score_cell_style
            )
            merge_info[(start_row, 1)] = (start_row, self.__max_score_col)
            merge_info[(start_row + 1, 1)] = (start_row + 1, self.__max_score_col)
            start_row += 2
            plus = 1
        table_info[start_row + plus, 0] = ReportTableCellData(value=data.name)
        for idx, benchmark in enumerate(data.benchmark_list, start=1):
            cur_idx = self.__max_indicator_col + idx - 1
            if first:
                table_info[start_row, cur_idx] = ReportTableCellData(
                    value=benchmark.tag, style=score_cell_style
                )
            table_info[start_row + plus, cur_idx] = ReportTableCellData(
                value=benchmark.score, style=score_cell_style
            )
        start_row = start_row + 2 if first else start_row + 1
        return start_row

    def traverse_root_data_in_level1(
        self,
        data: ReportIndicatorScoreViewModel,
        table_info: Dict,
        merge_info: Dict,
        first: bool,
        start_row,
    ) -> int:
        """
        转换根节点数据，对标签打在第一层的情况
        """
        if first:
            table_info[start_row, 0] = ReportTableCellData(
                value="评价内容",
                style=EnumReportStyle.EVALUATION_CONTENT_CELL.name,
            )
            table_info[start_row + 1, 0] = ReportTableCellData(
                value=data.name,
                style=EnumReportStyle.ROOT_INDICATOR_CELL.name,
            )
        else:
            table_info[start_row, 0] = ReportTableCellData(
                value=data.name,
                style=EnumReportStyle.ROOT_INDICATOR_CELL.name,
            )

        merge_info[(start_row, 0)] = (start_row, self.__max_indicator_col - 1)
        merge_info[(start_row + 1, 0)] = (start_row + 1, self.__max_indicator_col - 1)

        max_col = self.__max_score_col + self.__max_indicator_col - 1
        for idx, benchmark in enumerate(data.benchmark_list, start=1):
            cur_idx = self.__max_indicator_col + idx - 1
            table_info[start_row, cur_idx] = ReportTableCellData(value=benchmark.tag)
            table_info[start_row + 1, cur_idx] = ReportTableCellData(
                value=benchmark.score,
                root_score_cell=not first,
                style=EnumReportStyle.SCORE_CELL.name,
            )

            if idx == len(data.benchmark_list) and idx != self.__max_score_col:
                merge_info[(start_row, cur_idx)] = (start_row, max_col)
                merge_info[(start_row + 1, cur_idx)] = (start_row + 1, max_col)

        plus_row = 2
        tag_set = data.get_child_benchmark_tag(tag_set=set())
        if len(tag_set) == self.__max_score_col:
            # 如果这个指标的所有benchmark的tag长度和最大列一样，就显示tag，不然就不显示了
            # 如果从一个根开始这下面的benchmark的tag又有自评、他评又有计算类的综合，就不能画出来合适的样子，所以这里就不显示了
            # 用tag来作为自评、他评的文字显示，或许不是一个很好的选择，滨实先这样处理，产品的时候要仔细斟酌过
            for t_idx, tag in enumerate(tag_set, start=1):
                cur_idx = self.__max_indicator_col + t_idx - 1
                table_info[start_row + 2, cur_idx] = ReportTableCellData(value=tag)
            merge_info[(start_row + 2, 0)] = (start_row + 2, self.__max_indicator_col - 1)
            plus_row += 1
        return start_row + plus_row

    def traverse_child_data(
        self,
        data_list: List[ReportIndicatorScoreViewModel],
        table_data: Dict,
        merge_info: Dict,
        start_row,
        start_col,
    ):
        row = start_row
        col = start_col
        # 最大的列数
        max_col = self.__max_score_col + self.__max_indicator_col - 1
        for idx, data in enumerate(data_list, start=1):
            # 当前进行到的
            next_col = col
            table_data[row, col] = ReportTableCellData(value=data.name)
            for b_idx, benchmark in enumerate(data.benchmark_list, start=1):
                if b_idx > self.__max_score_col:
                    break
                cur_col = self.__max_indicator_col + b_idx - 1
                table_data[row, cur_col] = ReportTableCellData(
                    value=benchmark.score, style=EnumReportStyle.SCORE_CELL.name
                )
                if b_idx == len(data.benchmark_list) and b_idx != self.__max_score_col:
                    merge_info[(row, cur_col)] = (row, max_col)

            if col < self.__max_indicator_col - 1 and (not data.children or data.benchmark_list):
                # 如果这个单元格，后面是benchmark得分，或者这个格子没有孩子，那么需要向右合并
                merge_info[(row, col)] = (row, self.__max_indicator_col - 1)

            if data.benchmark_list:
                # 如果这一级有benchmark，孩子需要排在下一行
                row += 1
            else:
                # 如果这一级没有benchmark，孩子需要排在下一列
                next_col += 1

            if data.children:
                new_row = self.traverse_child_data(
                    data.children, table_data, merge_info, start_row=row, start_col=next_col
                )
                row = new_row
        return row

    def prepare_merge_data(self, table_info: Dict, merge_info: Dict):
        """
        准备需要合并的数据
        本来还有横向合并的，但是横向合并的太难判断了，换了另外一种写法，
        如果未来的某一天又有必要，可以参考2023-09-11的提交日志，我在这个版本先把横向合并的去掉了
        """
        self.merge_cell_vertically(table_info=table_info, merge_info=merge_info)

    def merge_cell_vertically(self, table_info: Dict, merge_info: Dict):
        """
        纵向合并单元格
        """
        max_row = max(row for row, _ in table_info.keys())
        for (row, col), row_data in table_info.items():
            if col > self.__max_indicator_col:
                # 纵向合并的目前只为了处理指标的问题
                continue
            merge_row = row
            count = 1
            while True:
                new_row = row + count
                # 向下合并单元格需要：1.新的行没有超过最大行数 2.新的单元格没有在table_info，说明这格没有内容填充
                # 3.新的行的第1个单元格也不应该是有内容填充的（主要对指标向右合并孩子指标的时候）
                if (
                    new_row <= max_row
                    and (new_row, col) not in table_info
                    and (new_row, 0) not in table_info
                ):
                    merge_row += 1
                    count += 1
                else:
                    break
            if merge_row > row:
                if (row, col) in merge_info:
                    # 如果这个单元格已经被合并过，可能是在横向合并的时候已经加进来了，在这里取最大的行号，列号进行合并
                    old_merge_data = merge_info[(row, col)]
                    merge_info[(row, col)] = (
                        max(merge_row, old_merge_data[0]),
                        max(col, old_merge_data[1]),
                    )
                else:
                    merge_info[(row, col)] = (merge_row, col)

    def create_table_in_word(
        self, doc: Document, table_info: Dict[Tuple, ReportTableCellData], merge_info: Dict
    ):
        # 计算表格的行数和列数
        total_rows = max(row for row, _ in table_info.keys()) + 1
        total_cols = max(col for _, col in table_info.keys()) + 1

        table = doc.add_table(rows=total_rows, cols=total_cols)
        table.style = "Table Grid"

        for (row, col), data in table_info.items():
            table.cell(row, col).text = data.value
            self.set_cell_style(cell=table.cell(row, col), style_type=data.style)

        for (start_row, start_col), (end_row, end_col) in merge_info.items():
            cell_start = table.cell(start_row, start_col)
            cell_end = table.cell(end_row, end_col)
            cell_start.merge(cell_end)

        self.adjust_table_width(table=table)
        table.autofit = False

    @staticmethod
    def set_cell_style(cell, style_type: str):
        """
        设置单元格字体
        :param cell: 单元格
        :param style_type: 样式类型
        :return:
        """
        tag_cell_style = ReportCellStyle(
            font_size=11.5,
            color=(255, 0, 0),
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        style_dict = {
            EnumReportStyle.TAG_CELL.name: tag_cell_style,
            EnumReportStyle.VERTICALLY_TAG_CELL.name: tag_cell_style,
            EnumReportStyle.SCORE_CELL.name: ReportCellStyle(
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            ),
            EnumReportStyle.ROOT_INDICATOR_CELL.name: ReportCellStyle(bold=True),
            EnumReportStyle.EVALUATION_CONTENT_CELL.name: ReportCellStyle(
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            ),
        }
        if style_type in [
            EnumReportStyle.VERTICALLY_TAG_CELL.name,
            EnumReportStyle.ROOT_INDICATOR_CELL.name,
            EnumReportStyle.EVALUATION_CONTENT_CELL.name,
        ]:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style = style_dict.get(style_type)
        if not style:
            style = ReportCellStyle()
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            run_obj = paragraph.runs
            for run in run_obj:
                font = run.font
                font.size = Pt(style.font_size)
                font.name = "仿宋"
                font.color.rgb = RGBColor(*style.color)
                font.bold = style.bold
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

            if style.alignment:
                paragraph.alignment = style.alignment

    def adjust_table_width(self, table: Table):
        """
        调整表格的宽度
        在office查看过word最合适的宽度是15.3cm
        """

        # 假设最小节点的宽度固定为7cm（如果理想状态和客户给的表格一直，这样的配置最合适，
        # 如果超出了预想的层级，最后一层内容较多还是先保证7cm，一般情况前面的文字少，可以挤一点竖着排列）
        max_table_width = 15.3
        last_indicator_col_width = 7
        col_widths = []

        if self.__max_indicator_col > 1:
            score_col_width = 1.7
            # 写的比较死，满足目前需要，这样写只是为了12345蔚蓝善行准则
            if self.__tag_level == 2 and self.__max_indicator_col == 3:
                col_widths.append(Cm(3.6))
                col_widths.append(Cm(6))
                col_widths.insert(0, Cm(15.3 - 3.6 - 6 - score_col_width * self.__max_score_col))
            else:
                other_indicator_width = (
                        max_table_width - score_col_width * self.__max_score_col - last_indicator_col_width
                )
                other_indicator_count = self.__max_indicator_col - 1
                if other_indicator_count:
                    col_widths = [Cm((other_indicator_width / other_indicator_count))] * other_indicator_count
                col_widths.append(Cm(last_indicator_col_width))
        else:
            score_col_width = 1.8
            col_widths.append(Cm(max_table_width - score_col_width * self.__max_score_col))

        for i in range(0, self.__max_score_col):
            col_widths.append(Cm(score_col_width))

        for col_idx, width in enumerate(col_widths):
            if width:  # 如果定义了宽度
                for cell in table.columns[col_idx].cells:
                    cell.width = width

    def add_report_table_for_tag_in_level2(self, document: Document):
        """
        添加报告的表格部分，针对标签打在第二层的情况
        """

        row = 0
        table_info: Dict[Tuple, ReportTableCellData] = {}
        merge_info = {}
        tag_tree = self.build_report_indicator_tree()
        # 因为X和N的标签竖着排列了，和标签打在第一层的相比，这种方式的时候X和N也是占的指标的列，所以需要加1
        self.__max_indicator_col += 1
        row = self.traverse_first_raw_in_level2(
            table_info=table_info,
            merge_info=merge_info,
            start_row=row,
        )
        for tag in ["N（常量）", "X（变量）"]:
            report_tree_data = tag_tree[tag]
            table_info[row, 0] = ReportTableCellData(
                value=tag, style=EnumReportStyle.VERTICALLY_TAG_CELL.name
            )

            row = self.traverse_child_data(report_tree_data, table_info, merge_info, row, 1)
        self.traverse_root_data_in_level2(
            table_info=table_info, merge_info=merge_info, start_row=row
        )
        self.merge_cell_vertically(table_info=table_info, merge_info=merge_info)
        self.create_table_in_word(doc=document, table_info=table_info, merge_info=merge_info)

    def traverse_first_raw_in_level2(
        self,
        table_info: Dict,
        merge_info: Dict,
        start_row,
    ) -> int:
        """
        转换第一行数据，对标签打在第一层的情况
        """
        table_info[start_row, 0] = ReportTableCellData(
            value="评价内容",
            style=EnumReportStyle.EVALUATION_CONTENT_CELL.name,
        )
        tag_list = self.__report_repository_v2.fetch_criteria_tree_child_benchmark_tag(
            evaluation_criteria_id=self.__plan.evaluation_criteria_id,
            executed_finish_at=self.__plan.executed_finish_at,
        )
        sorted_tag_list = sorted(tag_list, key=lambda x: ORDER.index(x))
        for t_idx, tag in enumerate(sorted_tag_list, start=1):
            cur_idx = self.__max_indicator_col + t_idx - 1
            table_info[start_row, cur_idx] = ReportTableCellData(value=tag, style=EnumReportStyle.SCORE_CELL.name)
        merge_info[(start_row, 0)] = (start_row, self.__max_indicator_col - 1)
        return start_row + 1

    def traverse_root_data_in_level2(self, table_info: Dict, merge_info: Dict, start_row):
        """
        转换标签打在第二层时的根节点的数据
        """

        data_list = self.__report_repository_v2.fetch_root_indicator_benchmark_score(
            evaluation_assignment_id=self.__plan.evaluation_assignment_id,
            evaluation_criteria_id=self.__plan.evaluation_criteria_id,
            executed_finish_at=self.__plan.executed_finish_at,
        )
        if len(data_list) == 1:
            # 追求速度，不考虑大于1的情况了，目前客户的需求只有1级的情况
            data = data_list[0]
            root_tag_list = ["综合", "等级"]
            for idx, root_tag in enumerate(root_tag_list):
                table_info[start_row + idx, 0] = ReportTableCellData(
                    value=root_tag,
                    style=EnumReportStyle.EVALUATION_CONTENT_CELL.name,  # 因为目前这个样式和评价内容的样式一样，偷个懒，本来应该重新定义一个的
                )
                merge_info[(start_row + idx, 0)] = (start_row + idx, self.__max_indicator_col - 1)
                # 直接根据名字匹配
                for benchmark in data.benchmark_list:
                    if benchmark.matched(root_tag, "自评"):
                        table_info[start_row + idx, self.__max_indicator_col] = ReportTableCellData(value=benchmark.score)
                    elif benchmark.matched(root_tag, "他评"):
                        table_info[start_row + idx, self.__max_indicator_col + 1] = ReportTableCellData(value=benchmark.score)
